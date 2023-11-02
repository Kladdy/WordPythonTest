import docx_svg_patch
import datetime

from docx import Document


SVG = """<svg version="1.1"
     width="300" height="200"
     xmlns="http://www.w3.org/2000/svg">

  <rect width="100%" height="100%" fill="red" />
  <circle cx="150" cy="100" r="80" fill="green" />
  <text x="150" y="125" font-size="60" text-anchor="middle" fill="white">SVG</text>
</svg>
"""
svg_filename = "drawing.svg"
with open(svg_filename, "w") as f:
    f.write(SVG)

document = Document()
now = datetime.datetime.now()
document.add_heading("Document Title", 0)
document.add_heading(f"{now}", 0)

document.add_picture(svg_filename)
document.add_picture("alphachannel.svg")

document.save("demo.docx")