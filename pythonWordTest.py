
# import docx_emf_patch
import docx_svg_patch
from docx import Document
import os


def main():
    template_file_path = 'template.docx'
    output_file_path = 'out.docx'

    variables = {
        "namn1": "lol",
        "namn2": "japp",
        "namn3": "tjenare",
        "namn4": "lmao",
    }

    template_document = Document(template_file_path)

    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, variable_key, variable_value)


    p = template_document.add_paragraph('Picture bullet section')
    p = p.insert_paragraph_before('')
    r = p.add_run()
    # r.add_picture("MLFicon.pdf")
    # r.add_picture("MLFicon.emf")
    # r.add_picture("MLFicon.png")
    r.add_picture("alphachannel.svg")

    template_document.save(output_file_path)


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


if __name__ == '__main__':
    main()